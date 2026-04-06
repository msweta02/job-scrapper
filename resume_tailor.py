import json
import re
import time
import requests
import pandas as pd
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Config ────────────────────────────────────────────────────────────────────

TOKEN_FILE    = "config.json"
MASTER_RESUME = "master_resume.docx"
OUTPUT_DIR    = Path("output")
GROQ_API_URL  = "https://api.groq.com/openai/v1/chat/completions"
GROQ_MODEL    = "llama-3.1-8b-instant"

# ── Load config ───────────────────────────────────────────────────────────────

def load_config(path: str) -> dict:
    with open(path, "r") as f:
        cfg = json.load(f)
    if not cfg.get("groq_api_key"):
        raise ValueError("'groq_api_key' not found in config.json")
    return cfg

# ── Extract text from master resume (.docx) ───────────────────────────────────

def extract_resume_text(docx_path: str) -> str:
    doc  = Document(docx_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if not text:
        raise ValueError(f"Master resume '{docx_path}' appears to be empty.")
    return text

# ── Build prompt ──────────────────────────────────────────────────────────────

def build_prompt(resume_text: str, job: dict) -> str:
    jd = job.get("description") or job.get("jobDescription") or "No description available."

    if len(jd) > 4000:
        jd = jd[:4000] + "\n... [truncated]"
    if len(resume_text) > 3000:
        resume_text = resume_text[:3000] + "\n... [truncated]"

    jd          = jd.encode("utf-8", errors="ignore").decode("utf-8")
    resume_text = resume_text.encode("utf-8", errors="ignore").decode("utf-8")

    return f"""
You are an expert resume writer and ATS optimization specialist.

## Master Resume (use as the ONLY source of facts):
{resume_text}

## Job Description:
Company:  {job.get('companyName') or job.get('company', 'N/A')}
Title:    {job.get('title') or job.get('jobTitle', 'N/A')}
Location: {job.get('location', 'N/A')}

{jd}

## Your Task:
Using ONLY facts from the master resume, rewrite a complete fresh ATS-optimized resume tailored for this job.

1. SUMMARY: Write a powerful 3-4 sentence summary tailored to this role.

2. EXPERIENCE: Rewrite bullet points for ALL 3 jobs in this exact order (latest to oldest):
   - JP Morgan Chase & Co (December 2017 – August 2024) → Senior Associate
   - EdgeVerve Systems Limited (August 2015 – December 2017) → Technology Analyst
   - Infosys Limited (June 2011 – July 2015) → Senior Software Engineer
   Write 3-4 strong bullets per role using ONLY facts from the master resume such as:
   - Technical: database migrations, SQL optimization, data pipelines, scripting, Spark, Airflow
   - Collaboration: client stakeholders, cross-functional teams, onsite/offshore
   - Leadership: mentoring junior engineers, knowledge transfer, team transitions
   - Quality: validation frameworks, data integrity, rollback mechanisms, resolving production issues
   NEVER write "No relevant information available". Always extract meaningful content for every role.

3. SKILLS: Reorder skills section putting most JD-relevant skills first.

4. EDUCATION: Include both degrees exactly as in the master resume.

5. CERTIFICATIONS: Include all certifications from the master resume.

6. SCORING:
   - relevance_score (0-100): How well the master resume experience matches this JD overall.
   - ats_score (0-100): Estimated ATS match score of the newly tailored resume vs this JD based on keyword overlap, skills match, and experience alignment.
   - gaps: Top 3-5 keywords/skills in JD missing from master resume.

## Output Format (return ONLY valid JSON, no markdown, no extra text):
{{
  "summary": "3-4 sentence tailored summary",
  "experience": [
    {{
      "company":  "Company name exactly as in master resume",
      "title":    "Job title exactly as in master resume",
      "location": "Location exactly as in master resume",
      "dates":    "Date range exactly as in master resume",
      "bullets":  [
        "Rewritten bullet 1 using STAR method",
        "Rewritten bullet 2",
        "Rewritten bullet 3",
        "Rewritten bullet 4"
      ]
    }}
  ],
  "skills": {{
    "Programming":      ["skill1", "skill2"],
    "Big Data":         ["skill1", "skill2"],
    "Cloud":            ["skill1", "skill2"],
    "Data Engineering": ["skill1", "skill2"],
    "Orchestration":    ["skill1", "skill2"],
    "DevOps":           ["skill1", "skill2"],
    "Databases":        ["skill1", "skill2"],
    "Other":            ["skill1", "skill2"]
  }},
  "education": [
    {{
      "institution": "University name",
      "degree":      "Degree name",
      "dates":       "Date range"
    }}
  ],
  "certifications": ["cert1", "cert2"],
  "relevance_score": 0,
  "ats_score":       0,
  "gaps":            ["missing1", "missing2"]
}}
"""

# ── Call Groq API ─────────────────────────────────────────────────────────────

def call_llm(api_key: str, prompt: str, retries: int = 5) -> dict:
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type":  "application/json",
    }
    body = {
        "model":       GROQ_MODEL,
        "messages":    [{"role": "user", "content": prompt}],
        "max_tokens":  3000,
        "temperature": 0.3,
    }

    for attempt in range(1, retries + 1):
        res = requests.post(GROQ_API_URL, headers=headers, json=body, timeout=60)

        if res.status_code == 429:
            retry_after = int(res.headers.get("retry-after", 0))
            wait        = retry_after if retry_after > 0 else (attempt * 15)
            print(f"\n         ⏳ Rate limited. Waiting {wait}s before retry {attempt}/{retries}...")
            time.sleep(wait)
            continue

        if not res.ok:
            try:
                err = res.json()
                msg = err.get("error", {}).get("message", res.text)
            except Exception:
                msg = res.text
            raise RuntimeError(f"Groq API {res.status_code}: {msg}")

        raw = res.json()["choices"][0]["message"]["content"].strip()

        # Log raw response for debugging if it looks empty or broken
        if not raw:
            print(f"\n         ⚠️  Empty response from Groq, retrying...")
            time.sleep(5)
            continue

        # Strip markdown fences if present
        raw = re.sub(r"^```json|^```|```$", "", raw, flags=re.MULTILINE).strip()

        # Extract JSON object if wrapped in extra text
        json_match = re.search(r"\{.*\}", raw, re.DOTALL)
        if not json_match:
            print(f"\n         ⚠️  No JSON found in response, retrying...")
            print(f"         Raw response: {raw[:200]}")
            time.sleep(5)
            continue

        return json.loads(json_match.group())

    raise RuntimeError(f"Groq rate limit exceeded after {retries} retries.")

# ── Write tailored resume as .docx ───────────────────────────────────────────

def write_docx(master_path: str, result: dict, job: dict, out_path: Path) -> None:
    doc     = Document()
    section = doc.sections[0]
    section.top_margin    = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin   = Pt(54)
    section.right_margin  = Pt(54)

    company = (job.get("companyName") or job.get("company") or "Company").strip()
    title   = (job.get("title")       or job.get("jobTitle") or "Role").strip()

    def add_heading(text: str) -> None:
        from docx.oxml.ns import qn
        from docx.oxml   import OxmlElement
        p   = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        pPr    = p._p.get_or_add_pPr()
        pBdr   = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F497D")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_bullet(text: str) -> None:
        p   = doc.add_paragraph()
        run = p.add_run("• " + text.lstrip("•").strip())
        run.font.size = Pt(10)
        p.paragraph_format.left_indent  = Pt(12)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

    # ── Header ────────────────────────────────────────────────────────────────
    name_para      = doc.add_paragraph()
    name_run       = name_para.add_run("Mary Sweta Kerketta")
    name_run.bold  = True
    name_run.font.size    = Pt(16)
    name_para.alignment   = WD_ALIGN_PARAGRAPH.CENTER

    contact_para = doc.add_paragraph()
    contact_run  = contact_para.add_run(
        "Boston, MA  |  857-379-6775  |  kerketta.m@northeastern.edu  |  linkedin.com/in/marysweta"
    )
    contact_run.font.size           = Pt(10)
    contact_para.alignment          = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(4)

    # ── Summary ───────────────────────────────────────────────────────────────
    add_heading("Summary")
    summary_p = doc.add_paragraph(result.get("summary", ""))
    summary_p.runs[0].font.size           = Pt(10)
    summary_p.paragraph_format.space_after = Pt(4)

    # ── Experience ────────────────────────────────────────────────────────────
    add_heading("Experience")
    for exp in result.get("experience", []):
        exp_header = doc.add_paragraph()
        co_run     = exp_header.add_run(exp.get("company", ""))
        co_run.bold      = True
        co_run.font.size = Pt(10)
        date_run         = exp_header.add_run(f"    {exp.get('dates', '')}")
        date_run.font.size                    = Pt(10)
        exp_header.paragraph_format.space_before = Pt(6)
        exp_header.paragraph_format.space_after  = Pt(1)

        role_p   = doc.add_paragraph()
        role_run = role_p.add_run(
            f"{exp.get('title', '')}  |  {exp.get('location', '')}"
        )
        role_run.italic    = True
        role_run.font.size = Pt(10)
        role_p.paragraph_format.space_after = Pt(2)

        for bullet in exp.get("bullets", []):
            add_bullet(bullet)

    # ── Skills ────────────────────────────────────────────────────────────────
    add_heading("Technical Skills")
    for category, items in result.get("skills", {}).items():
        if not items:
            continue
        p        = doc.add_paragraph()
        cat_run  = p.add_run(f"{category}: ")
        cat_run.bold      = True
        cat_run.font.size = Pt(10)
        items_run         = p.add_run(", ".join(items))
        items_run.font.size             = Pt(10)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    # ── Education ─────────────────────────────────────────────────────────────
    add_heading("Education")
    for edu in result.get("education", []):
        edu_p    = doc.add_paragraph()
        edu_run  = edu_p.add_run(edu.get("institution", ""))
        edu_run.bold      = True
        edu_run.font.size = Pt(10)
        date_run          = edu_p.add_run(f"    {edu.get('dates', '')}")
        date_run.font.size                   = Pt(10)
        edu_p.paragraph_format.space_before  = Pt(4)
        edu_p.paragraph_format.space_after   = Pt(1)

        deg_p   = doc.add_paragraph()
        deg_run = deg_p.add_run(edu.get("degree", ""))
        deg_run.italic    = True
        deg_run.font.size = Pt(10)

    # ── Certifications ────────────────────────────────────────────────────────
    certs = result.get("certifications", [])
    if certs:
        add_heading("Certifications")
        for cert in certs:
            add_bullet(cert)

    doc.save(out_path)

# ── Safe filename ─────────────────────────────────────────────────────────────

def safe_filename(company: str, title: str) -> str:
    raw = f"{company}_{title}"
    return re.sub(r"[^\w\-]", "_", raw)[:80]

# ── Save report ───────────────────────────────────────────────────────────────

def save_report(report_rows: list[dict], daily_dir: Path) -> None:
    if not report_rows:
        return

    total    = len(report_rows)
    avg_rel  = sum(r["relevance_score"] for r in report_rows) / total
    avg_ats  = sum(r["ats_score"]       for r in report_rows) / total

    # ── Build dataframe ───────────────────────────────────────────────────────
    rows = []
    for i, r in enumerate(report_rows, 1):
        rows.append({
            "#":               i,
            "Company":         r["company"],
            "Job Title":       r["job_title"],
            "Status":          r["status"].capitalize(),
            "Resume File":     r["filename"],
            "Relevance Score": r["relevance_score"],
            "ATS Score":       r["ats_score"],
            "Top Keywords":    r["keywords"],
            "Gaps":            r["gaps"],
            "Error":           r["error"],
        })

    df = pd.DataFrame(rows)
    df.sort_values(
        by=["Relevance Score", "ATS Score", "Company"],
        ascending=[False, False, True],
        inplace=True,
        na_position="last",
    )
    df["#"] = range(1, len(df) + 1)  # re-number after sort

    report_path = daily_dir / "resume_tailor_report.xlsx"

    with pd.ExcelWriter(report_path, engine="openpyxl") as writer:
        # ── Jobs sheet ────────────────────────────────────────────────────────
        df.to_excel(writer, index=False, sheet_name="Jobs")
        ws = writer.sheets["Jobs"]

        from openpyxl.styles import Alignment, PatternFill, Font, Border, Side
        from openpyxl.utils  import get_column_letter

        header_fill   = PatternFill("solid", fgColor="1F497D")
        header_font   = Font(color="FFFFFF", bold=True, size=10)
        success_fill  = PatternFill("solid", fgColor="E2EFDA")  # green
        failed_fill   = PatternFill("solid", fgColor="FFDDC1")  # orange
        skipped_fill  = PatternFill("solid", fgColor="FFF2CC")  # yellow
        thin_border   = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin")
        )

        # Style header row
        for cell in ws[1]:
            cell.fill      = header_fill
            cell.font      = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border    = thin_border

        # Style data rows
        for row in ws.iter_rows(min_row=2):
            status = str(row[3].value or "").lower()
            fill   = success_fill if status == "success" else (
                     failed_fill  if status == "failed"  else skipped_fill)
            for cell in row:
                cell.fill      = fill
                cell.border    = thin_border
                cell.font      = Font(size=10)
                cell.alignment = Alignment(vertical="top", wrap_text=True)

        # Color-code score cells (Relevance=col F, ATS=col G → indices 5,6)
        for row in ws.iter_rows(min_row=2):
            for col_idx in [5, 6]:
                cell  = row[col_idx]
                score = cell.value or 0
                if score >= 75:
                    cell.font = Font(color="375623", bold=True, size=10)  # dark green
                elif score >= 50:
                    cell.font = Font(color="7F6000", bold=True, size=10)  # dark yellow
                else:
                    cell.font = Font(color="9C0006", bold=True, size=10)  # dark red
                cell.alignment = Alignment(horizontal="center", vertical="top")

        # Auto-fit column widths
        col_widths = {"#": 5, "Company": 25, "Job Title": 35, "Status": 10,
                      "Resume File": 40, "Relevance Score": 10, "ATS Score": 10,
                      "Top Keywords": 35, "Gaps": 30, "Error": 30}
        for i, col in enumerate(df.columns, 1):
            ws.column_dimensions[get_column_letter(i)].width = col_widths.get(col, 20)

        # Freeze header row
        ws.freeze_panes = "A2"

        # ── Summary sheet ─────────────────────────────────────────────────────
        ws2 = writer.book.create_sheet("Summary")
        summary_data = [
            ("Generated",            datetime.now().strftime("%Y-%m-%d %H:%M")),
            ("Model",                GROQ_MODEL),
            ("Total Jobs",           total),
            ("Successful",           sum(1 for r in report_rows if r["status"] == "success")),
            ("Skipped (no JD)",      sum(1 for r in report_rows if r["status"] == "skipped")),
            ("Failed",               sum(1 for r in report_rows if r["status"] == "failed")),
            ("Avg Relevance Score",  f"{avg_rel:.1f}/100"),
            ("Avg ATS Score",        f"{avg_ats:.1f}/100"),
        ]
        for row_idx, (label, value) in enumerate(summary_data, 1):
            ws2.cell(row=row_idx, column=1, value=label).font  = Font(bold=True, size=10)
            ws2.cell(row=row_idx, column=2, value=str(value)).font = Font(size=10)
        ws2.column_dimensions["A"].width = 25
        ws2.column_dimensions["B"].width = 25

    print(f"\n📊 Report saved → {report_path.resolve()}")

# ── Main tailor function ──────────────────────────────────────────────────────

def tailor_resumes(excel_path: Path, resumes_dir: Path) -> None:
    cfg        = load_config(TOKEN_FILE)
    api_key    = cfg["groq_api_key"]
    resume_txt = extract_resume_text(MASTER_RESUME)
    daily_dir  = excel_path.parent

    df         = pd.read_excel(excel_path)
    jobs       = df.to_dict(orient="records")
    report_rows = []

    print(f"\n🤖 Tailoring resumes for {len(jobs)} jobs...")
    print(f"   Model  → {GROQ_MODEL}")
    print(f"   Output → {resumes_dir.resolve()}\n")

    for i, job in enumerate(jobs, 1):
        company = (str(job.get("Company")         or "")).strip()
        title   = (str(job.get("Job Title")       or "")).strip()
        jd      =  str(job.get("Job Description") or "")

        if not jd or jd.lower() in ("nan", "none", ""):
            print(f"  [{i}/{len(jobs)}] ⚠️  Skipping {company} — no JD available.")
            report_rows.append({
                "company": company, "job_title": title,
                "status": "skipped", "filename": "",
                "relevance_score": 0, "ats_score": 0,
                "keywords": "", "gaps": "", "error": "No JD available",
            })
            continue

        print(f"  [{i}/{len(jobs)}] ✍️  {company} — {title}")

        try:
            job_dict = {
                "companyName": company,
                "title":       title,
                "location":    str(job.get("Location") or ""),
                "description": jd,
            }

            prompt   = build_prompt(resume_txt, job_dict)
            result   = call_llm(api_key, prompt)

            fname    = safe_filename(company, title)
            out_path = resumes_dir / f"{fname}.docx"
            write_docx(MASTER_RESUME, result, job_dict, out_path)

            rel_score = result.get("relevance_score", 0)
            ats_score = result.get("ats_score",       0)
            keywords  = ", ".join(result.get("top_keywords", [])[:5])
            gaps      = ", ".join(result.get("gaps",          [])[:5])

            print(f"         💾 Saved            : {out_path.name}")
            print(f"         📊 Relevance Score  : {rel_score}/100")
            print(f"         🎯 ATS Score        : {ats_score}/100")
            print(f"         🔑 Top Keywords     : {keywords or '—'}")
            print(f"         ⚠️  Gaps             : {gaps     or 'none'}")

            report_rows.append({
                "company": company, "job_title": title,
                "status": "success", "filename": out_path.name,
                "relevance_score": rel_score, "ats_score": ats_score,
                "keywords": keywords, "gaps": gaps, "error": "",
            })

        except json.JSONDecodeError as e:
            print(f"         ❌ JSON parse error : {e}")
            print(f"         ℹ️  This usually means Groq returned incomplete JSON due to token limits.")
            print(f"         ℹ️  Try reducing JD length or increasing time.sleep between jobs.")
            report_rows.append({
                "company": company, "job_title": title,
                "status": "failed", "filename": "",
                "relevance_score": 0, "ats_score": 0,
                "keywords": "", "gaps": "", "error": str(e),
            })
        except RuntimeError as e:
            print(f"         ❌ Failed            : {e}")
            report_rows.append({
                "company": company, "job_title": title,
                "status": "failed", "filename": "",
                "relevance_score": 0, "ats_score": 0,
                "keywords": "", "gaps": "", "error": str(e),
            })
        except Exception as e:
            print(f"         ❌ Unexpected error  : {e}")
            report_rows.append({
                "company": company, "job_title": title,
                "status": "failed", "filename": "",
                "relevance_score": 0, "ats_score": 0,
                "keywords": "", "gaps": "", "error": str(e),
            })

        time.sleep(20)

    save_report(report_rows, daily_dir)
    print(f"\n✅ Resume tailoring complete → {resumes_dir.resolve()}")


# ── Standalone entry point ────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        folders = sorted(OUTPUT_DIR.glob("????-??-??"), reverse=True)
        if not folders:
            print("❌ No output folders found. Run scraper.py first.")
            sys.exit(1)
        excels = sorted(folders[0].glob("*.xlsx"), reverse=True)
        if not excels:
            print(f"❌ No Excel file found in {folders[0]}")
            sys.exit(1)
        excel_path = excels[0]
    else:
        excel_path = Path(sys.argv[1])

    if not excel_path.exists():
        print(f"❌ File not found: {excel_path}")
        sys.exit(1)

    resumes_dir = excel_path.parent / "resumes"
    resumes_dir.mkdir(exist_ok=True)

    print(f"📄 Using Excel : {excel_path}")
    print(f"📁 Resumes dir : {resumes_dir}")

    tailor_resumes(excel_path, resumes_dir)